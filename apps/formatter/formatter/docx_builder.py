from __future__ import annotations

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from formatter.config import FormatConfig


def _set_style_fonts(style, ascii_font: str, east_asia_font: str | None = None) -> None:
    style.font.name = ascii_font
    rfonts = style.element.rPr.rFonts
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    if east_asia_font:
        rfonts.set(qn("w:eastAsia"), east_asia_font)


def _lines_to_pt(lines: float, font_size_pt: int) -> Pt:
    return Pt(lines * font_size_pt)


def _chars_to_pt(chars: int, font_size_pt: int) -> Pt:
    return Pt(chars * font_size_pt)


def _apply_style_paragraph(style, size_pt, line_spacing, before_lines, after_lines, align=None) -> None:
    style.font.size = Pt(size_pt)
    pf = style.paragraph_format
    pf.space_before = _lines_to_pt(before_lines, size_pt)
    pf.space_after = _lines_to_pt(after_lines, size_pt)
    pf.line_spacing = line_spacing
    if align is not None:
        pf.alignment = align


def _add_page_number(section, position: str) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if position == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def _add_runs(paragraph, runs: list[dict], fallback_text: str = "") -> None:
    if not runs:
        if fallback_text:
            paragraph.add_run(fallback_text)
        return
    for run in runs:
        text = run.get("text", "")
        if not text:
            continue
        docx_run = paragraph.add_run(text)
        if run.get("bold"):
            docx_run.bold = True


def build_docx(ast: list[dict], output_path, config: FormatConfig | None = None) -> None:
    config = config or FormatConfig()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    try:
        _add_page_number(section, config.page_num_position)
    except Exception:
        pass

    try:
        normal = doc.styles["Normal"]
        _set_style_fonts(normal, config.body_style.en_font, config.body_style.cn_font)
        _apply_style_paragraph(
            normal,
            config.body_style.size_pt,
            config.body_style.line_spacing,
            config.body_style.para_before_lines,
            config.body_style.para_after_lines,
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY if config.body_style.justify else None,
        )
        normal.paragraph_format.left_indent = _chars_to_pt(
            config.body_style.indent_before_chars, config.body_style.size_pt
        )
        normal.paragraph_format.right_indent = _chars_to_pt(
            config.body_style.indent_after_chars, config.body_style.size_pt
        )
        normal.paragraph_format.first_line_indent = _chars_to_pt(
            config.body_style.first_line_indent_chars, config.body_style.size_pt
        )

        for level, hstyle in config.heading_styles.items():
            h = doc.styles[f"Heading {level}"]
            _set_style_fonts(h, hstyle.font, hstyle.font)
            h.font.color.rgb = RGBColor(0, 0, 0)
            _apply_style_paragraph(
                h,
                hstyle.size_pt,
                hstyle.line_spacing,
                hstyle.para_before_lines,
                hstyle.para_after_lines,
            )
    except Exception:
        pass

    for node in ast:
        if node.get("type") == "heading":
            paragraph = doc.add_heading("", level=node.get("level", 1))
            _add_runs(paragraph, node.get("runs", []), node.get("text", ""))
        elif node.get("type") == "paragraph":
            paragraph = doc.add_paragraph("")
            _add_runs(paragraph, node.get("runs", []), node.get("text", ""))
            if config.body_style.justify:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.left_indent = _chars_to_pt(
                config.body_style.indent_before_chars, config.body_style.size_pt
            )
            paragraph.paragraph_format.right_indent = _chars_to_pt(
                config.body_style.indent_after_chars, config.body_style.size_pt
            )
            paragraph.paragraph_format.first_line_indent = _chars_to_pt(
                config.body_style.first_line_indent_chars, config.body_style.size_pt
            )

    doc.save(output_path)
